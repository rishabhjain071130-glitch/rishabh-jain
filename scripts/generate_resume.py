import os

# Create scripts directory if it doesn't exist
os.makedirs(os.path.dirname(__file__), exist_ok=True)

# Generate DOCX and PDF resume


def main():
    print("Generating Resume_v2.docx...")
    generate_docx("Resume_v2.docx")
    print("Generating Resume_v2.pdf...")
    generate_pdf("Resume_v2.pdf")

    # Copy the generated PDF to public/resume/Resume_v2.pdf
    import shutil
    os.makedirs("public/resume", exist_ok=True)
    shutil.copyfile("Resume_v2.pdf", "public/resume/Resume_v2.pdf")

    print("Generation complete!")


def generate_docx(output_path):
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn

    doc = docx.Document()

    # Page setup - 0.5 in margins for one-page fit
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)

    # Style configuration
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(9.5)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)  # Off-black
    normal_style.paragraph_format.line_spacing = 1.05
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.space_before = Pt(0)

    # Helper function for horizontal rules
    def add_p_border_bottom(paragraph):
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                         r'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="A0A0A0"/>'
                         r'</w:pBdr>')
        pPr.append(pBdr)

    # Helper for adding sections
    def add_section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)  # Deep Navy
        add_p_border_bottom(p)
        return p

    def add_hyperlink(paragraph, text, url,
                      color_rgb=(0x1B, 0x36, 0x5D), underline=True,
                      font_size=Pt(8.5), bold=False, italic=False):
        """
        Adds a clickable hyperlink run to a paragraph.
        """
        part = paragraph.part
        r_id = part.relate_to(
            url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
        rPr.append(rFonts)

        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_size.pt * 2)))
            rPr.append(sz)

        if color_rgb:
            color_hex = f"{color_rgb[0]:02X}{color_rgb[1]:02X}{color_rgb[2]:02X}"
            c = OxmlElement('w:color')
            c.set(qn('w:val'), color_hex)
            rPr.append(c)

        if underline:
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)

        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)

        if italic:
            i = OxmlElement('w:i')
            rPr.append(i)

        new_run.append(rPr)

        text_node = OxmlElement('w:t')
        text_node.text = text
        new_run.append(text_node)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink

    # --- HEADER ---
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.paragraph_format.space_after = Pt(1)
    name_run = header_p.add_run("RISHABH JAIN\n")
    name_run.bold = True
    name_run.font.size = Pt(16)
    name_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(4)

    def add_text_run(p, text):
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        return run

    add_text_run(contact_p, "+91 9258121291  |  ")
    add_hyperlink(contact_p, "rishabhjain071130@gmail.com",
                  "mailto:rishabhjain071130@gmail.com",
                  color_rgb=(0x1B, 0x36, 0x5D), underline=True)
    add_text_run(contact_p, "  |  ")
    add_hyperlink(contact_p, "rishabh-jain-one.vercel.app",
                  "https://rishabh-jain-one.vercel.app",
                  color_rgb=(0x1B, 0x36, 0x5D), underline=True)
    add_text_run(contact_p, "\n")
    add_hyperlink(contact_p, "github.com/rishabhjain071130-glitch/rishabh-jain",
                  "https://github.com/rishabhjain071130-glitch/rishabh-jain",
                  color_rgb=(0x1B, 0x36, 0x5D), underline=True)
    add_text_run(contact_p, "  |  ")
    add_hyperlink(contact_p, "linkedin.com/in/rishabh-jain-40079a396",
                  "https://www.linkedin.com/in/rishabh-jain-40079a396/",
                  color_rgb=(0x1B, 0x36, 0x5D), underline=True)

    # --- 1. PROFESSIONAL SUMMARY ---
    add_section_header("Professional Summary")
    summary_p = doc.add_paragraph(
        "Computer Science Student specializing in Full Stack Development, "
        "Artificial Intelligence, and Cyber Security. "
        "Proven track record of engineering secure, intelligent, and "
        "scalable software through hands-on, real-world projects. "
        "Equipped with strong analytical problem-solving skills and "
        "industry-standard technical certifications."
    )
    summary_p.paragraph_format.space_after = Pt(3)

    # --- 2. EDUCATION ---
    add_section_header("Education")

    edu_table = doc.add_table(rows=3, cols=2)
    edu_table.autofit = False
    edu_table.columns[0].width = Inches(6.0)
    edu_table.columns[1].width = Inches(1.5)

    # GLA University
    cell_gla_left = edu_table.cell(0, 0)
    cell_gla_left.text = ""
    p_gla = cell_gla_left.paragraphs[0]
    r_gla_bold = p_gla.add_run("GLA University")
    r_gla_bold.bold = True
    p_gla.add_run(
        ", Diploma in Computer Science and Engineering (Lateral Entry)")

    cell_gla_right = edu_table.cell(0, 1)
    p_gla_r = cell_gla_right.paragraphs[0]
    p_gla_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_gla_r.add_run("Expected 2027")

    # SMS Jain Class 12
    cell_sms12_left = edu_table.cell(1, 0)
    p_sms12 = cell_sms12_left.paragraphs[0]
    r_sms12_bold = p_sms12.add_run("S.M.S Jain Inter College (UP Board)")
    r_sms12_bold.bold = True
    p_sms12.add_run(", Intermediate (Class 12) – Score: 78.2%")

    cell_sms12_right = edu_table.cell(1, 1)
    p_sms12_r = cell_sms12_right.paragraphs[0]
    p_sms12_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sms12_r.add_run("2024")

    # SMS Jain Class 10
    cell_sms10_left = edu_table.cell(2, 0)
    p_sms10 = cell_sms10_left.paragraphs[0]
    r_sms10_bold = p_sms10.add_run("S.M.S Jain Inter College (UP Board)")
    r_sms10_bold.bold = True
    p_sms10.add_run(", High School (Class 10) – Score: 77.0%")

    cell_sms10_right = edu_table.cell(2, 1)
    p_sms10_r = cell_sms10_right.paragraphs[0]
    p_sms10_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sms10_r.add_run("2023")

    for row in edu_table.rows:
        row.height = Pt(13)
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(1)
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.width = cell.width

    # --- 3. TECHNICAL SKILLS ---
    add_section_header("Technical Skills")
    skills_data = [
        ("Programming", "JavaScript, TypeScript, Python, HTML, CSS"),
        ("Frameworks", "React, Next.js, Node.js, Tailwind CSS"),
        ("Databases", "MongoDB, MySQL"),
        ("Tools", "Git, GitHub, Vercel, EmailJS, VS Code, Linux")
    ]
    for category, list_s in skills_data:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run_cat = p.add_run(f"{category}: ")
        run_cat.bold = True
        p.add_run(list_s)

    # --- 4. EXPERIENCE ---
    add_section_header("Experience")

    # Labmentix
    exp1_table = doc.add_table(rows=1, cols=2)
    exp1_table.columns[0].width = Inches(5.5)
    exp1_table.columns[1].width = Inches(2.0)
    c1 = exp1_table.cell(0, 0).paragraphs[0]
    c1_bold = c1.add_run("Labmentix")
    c1_bold.bold = True
    c1.add_run(" – Web Development Intern")
    c1.paragraph_format.space_after = Pt(2)

    c2 = exp1_table.cell(0, 1).paragraphs[0]
    c2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    c2.add_run("Feb 2026 – Present")
    c2.paragraph_format.space_after = Pt(2)

    bullets_labmentix = [
        "Built responsive React frontend components, improving UI performance and cross-device compatibility.",
        "Integrated Express backend controllers with MongoDB databases to implement secure web endpoints.",
        "Authored comprehensive API documentation and local setup guides to accelerate developer onboarding.",
        "Participated in peer code reviews, troubleshooting codebase bugs and improving code readability."
    ]
    for b in bullets_labmentix:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.2)

    # DecodeLabs
    exp2_table = doc.add_table(rows=1, cols=2)
    exp2_table.columns[0].width = Inches(5.5)
    exp2_table.columns[1].width = Inches(2.0)
    c1 = exp2_table.cell(0, 0).paragraphs[0]
    c1_bold = c1.add_run("DecodeLabs")
    c1_bold.bold = True
    c1.add_run(" – Web Development Intern")
    c1.paragraph_format.space_after = Pt(2)

    c2 = exp2_table.cell(0, 1).paragraphs[0]
    c2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    c2.add_run("Summer 2025")
    c2.paragraph_format.space_after = Pt(2)

    bullets_decodelabs = [
        "Created responsive, mobile-friendly layouts and modular site sections using React and Tailwind CSS.",
        "Connected React client components to Node.js REST API endpoints to deliver dynamic data updates.",
        "Utilized Git and GitHub for version control, branching, and pull requests in a team environment.",
        "Conducted local component and route testing to ensure strict alignment with visual design mockups."
    ]
    for b in bullets_decodelabs:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.2)

    # --- 5. FEATURED PROJECTS ---
    add_section_header("Featured Projects")

    # Project 1: Portfolio OS
    p_os = doc.add_paragraph()
    r_os = p_os.add_run("Rishabh Portfolio OS")
    r_os.bold = True
    p_os.add_run(" (Next.js, React, TypeScript, Tailwind CSS, Git) | ")
    add_hyperlink(
        p_os,
        "github.com/rishabhjain071130-glitch/rishabh-jain",
        "https://github.com/rishabhjain071130-glitch/rishabh-jain",
        color_rgb=(0x1B, 0x36, 0x5D), underline=True,
        font_size=Pt(9.5), italic=True
    )
    p_os.paragraph_format.space_after = Pt(1)
    p_os.paragraph_format.keep_with_next = True

    p = doc.add_paragraph(
        "Designed terminal-themed developer portfolio featuring CLI "
        "console simulation and rate-limited forms.",
        style='List Bullet'
    )
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.2)

    # Project 2: EduPilot AI
    p_ep = doc.add_paragraph()
    r_ep = p_ep.add_run("EduPilot AI")
    r_ep.bold = True
    p_ep.add_run(" (Python, Streamlit, Gemini API, HTML, CSS, Git) | ")
    add_hyperlink(
        p_ep,
        "github.com/rishabhjain071130-glitch/EduPilot-AI",
        "https://github.com/rishabhjain071130-glitch/EduPilot-AI",
        color_rgb=(0x1B, 0x36, 0x5D), underline=True,
        font_size=Pt(9.5), italic=True
    )
    p_ep.paragraph_format.space_after = Pt(1)
    p_ep.paragraph_format.keep_with_next = True

    p = doc.add_paragraph(
        "Built AI career mentor mapping skill gaps and generating "
        "custom roadmaps using Gemini API prompts.",
        style='List Bullet'
    )
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.2)

    # Project 3: Digital Loan System
    p_dl = doc.add_paragraph()
    r_dl = p_dl.add_run("Digital Loan System")
    r_dl.bold = True
    p_dl.add_run(" (Node.js, Express.js, MongoDB, Mongoose, HTML, CSS) | ")
    add_hyperlink(
        p_dl,
        "github.com/rishabhjain071130-glitch/digitalloansystem",
        "https://github.com/rishabhjain071130-glitch/digitalloansystem",
        color_rgb=(0x1B, 0x36, 0x5D), underline=True,
        font_size=Pt(9.5), italic=True
    )
    p_dl.paragraph_format.space_after = Pt(1)
    p_dl.paragraph_format.keep_with_next = True

    p = doc.add_paragraph(
        "Developed Express/MongoDB loan manager with session control, ledger logs, and automated notifications.",
        style='List Bullet'
    )
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.2)

    # --- 6. CERTIFICATIONS ---
    add_section_header("Certifications")
    certs = [
        "**Oracle Database SQL Certified Associate** – SQL queries, "
        "database indexing, and relational models (Dec 2025)",
        "**Microsoft Agentic AI Certification** – Development and deployment "
        "of autonomous LLM agent systems (June 2026)",
        "**Microsoft Web Technology Specialization** – Standard-compliant "
        "frontend and backend implementations (June 2026)",
        "**Palo Alto Networks Network Security Gateway** – Network defense "
        "and firewall rule configurations (June 2026)",
        "**IBM Machine Learning Roadmap Certification** – Core ML models, "
        "pipelines, and data architectures (Jan 2026)",
        "**Google Generative AI Course Completion** – Prompt engineering, "
        "LLM tuning, and AI safety practices (May 2026)",
        "**Internshala Cyber Security with AI Training** – Threat detection "
        "and AI-assisted security response (Jan 2026)"
    ]
    for c in certs:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.2)
        parts = c.split("**")
        if len(parts) >= 3:
            p.add_run(parts[1]).bold = True
            p.add_run(parts[2])
        else:
            p.add_run(c)

    # --- 7. ACHIEVEMENTS ---
    add_section_header("Achievements")
    achievements = [
        "Earned 10+ professional technical credentials from Microsoft, Oracle, Palo Alto Networks, and IBM.",
        "Deployed 3 production-grade open-source applications addressing real-world utility and security.",
        "Maintained a strong academic record and verified internship credentials in GLA University cohort."
    ]
    for a in achievements:
        p = doc.add_paragraph(a, style='List Bullet')
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.2)

    doc.save(output_path)


def generate_pdf(output_path):
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT

    # Create document with 0.5-inch margins (36 points) top/bottom margins reduced to 30 for safe one-page fit
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        leftMargin=36,
        rightMargin=36,
        topMargin=30,
        bottomMargin=30
    )

    story = []

    # Palette definition
    navy = HexColor("#1B365D")
    dark_grey = HexColor("#333333")
    light_grey = HexColor("#555555")
    rule_grey = HexColor("#A0A0A0")

    # Font styles
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'HeaderTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=16,
        leading=18,
        alignment=TA_CENTER,
        textColor=navy
    )

    contact_style = ParagraphStyle(
        'HeaderContact',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.2,
        leading=10.5,
        alignment=TA_CENTER,
        textColor=light_grey
    )

    section_title_style = ParagraphStyle(
        'SectionTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9.8,
        leading=11,
        textColor=navy,
        spaceBefore=4,
        spaceAfter=1,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'BodyText',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.8,
        leading=11,
        textColor=dark_grey,
        spaceAfter=2
    )

    bullet_style = ParagraphStyle(
        'BulletText',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.5,
        leading=10.5,
        textColor=dark_grey,
        leftIndent=10,
        firstLineIndent=-10,
        spaceAfter=0.8
    )

    def add_section_header(title_text):
        story.append(Paragraph(title_text.upper(), section_title_style))
        story.append(HRFlowable(width="100%", thickness=0.8,
                     color=rule_grey, spaceBefore=0.5, spaceAfter=2.5))

    # --- HEADER ---
    story.append(Paragraph("RISHABH JAIN", title_style))
    story.append(Spacer(1, 1.5))
    story.append(Paragraph(
        "+91 9258121291  |  "
        "<a href=\"mailto:rishabhjain071130@gmail.com\"><font color=\"#1B365D\">"
        "<u>rishabhjain071130@gmail.com</u></font></a>  |  "
        "<a href=\"https://rishabh-jain-one.vercel.app\"><font color=\"#1B365D\">"
        "<u>rishabh-jain-one.vercel.app</u></font></a><br/>"
        "<a href=\"https://github.com/rishabhjain071130-glitch/rishabh-jain\">"
        "<font color=\"#1B365D\"><u>github.com/rishabhjain071130-glitch/rishabh-jain</u>"
        "</font></a>  |  "
        "<a href=\"https://www.linkedin.com/in/rishabh-jain-40079a396/\">"
        "<font color=\"#1B365D\"><u>linkedin.com/in/rishabh-jain-40079a396</u>"
        "</font></a>",
        contact_style
    ))
    story.append(Spacer(1, 3))

    # --- 1. PROFESSIONAL SUMMARY ---
    add_section_header("Professional Summary")
    summary_text = (
        "Computer Science Student specializing in Full Stack Development, "
        "Artificial Intelligence, and Cyber Security. "
        "Proven track record of engineering secure, intelligent, and "
        "scalable software through hands-on, real-world projects. "
        "Equipped with strong analytical problem-solving skills and "
        "industry-standard technical certifications."
    )
    story.append(Paragraph(summary_text, body_style))
    story.append(Spacer(1, 1))

    # --- 2. EDUCATION ---
    add_section_header("Education")

    edu_data = [
        [
            Paragraph(
                "<b>GLA University</b>, Diploma in Computer Science and "
                "Engineering (Lateral Entry)", body_style
            ),
            Paragraph("Expected 2027", ParagraphStyle(
                'RightText', parent=body_style, alignment=TA_RIGHT))
        ],
        [
            Paragraph(
                "<b>S.M.S Jain Inter College (UP Board)</b>, Intermediate (Class 12) – Score: 78.2%", body_style),
            Paragraph("2024", ParagraphStyle(
                'RightText', parent=body_style, alignment=TA_RIGHT))
        ],
        [
            Paragraph(
                "<b>S.M.S Jain Inter College (UP Board)</b>, High School (Class 10) – Score: 77.0%", body_style),
            Paragraph("2023", ParagraphStyle(
                'RightText', parent=body_style, alignment=TA_RIGHT))
        ]
    ]
    edu_table = Table(edu_data, colWidths=[420, 120])
    edu_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0.5),
        ('TOPPADDING', (0, 0), (-1, -1), 0.5),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(edu_table)
    story.append(Spacer(1, 1))

    # --- 3. TECHNICAL SKILLS ---
    add_section_header("Technical Skills")
    skills_list = [
        ("Programming", "JavaScript, TypeScript, Python, HTML, CSS"),
        ("Frameworks", "React, Next.js, Node.js, Tailwind CSS"),
        ("Databases", "MongoDB, MySQL"),
        ("Tools", "Git, GitHub, Vercel, EmailJS, VS Code, Linux")
    ]
    for category, skills in skills_list:
        text = f"<b>{category}:</b> {skills}"
        story.append(Paragraph(text, body_style))
    story.append(Spacer(1, 1))

    # --- 4. EXPERIENCE ---
    add_section_header("Experience")

    # Labmentix
    labmentix_header = [
        [
            Paragraph("<b>Labmentix</b> – Web Development Intern", body_style),
            Paragraph("Feb 2026 – Present", ParagraphStyle(
                'RightText', parent=body_style, alignment=TA_RIGHT))
        ]
    ]
    t_lab = Table(labmentix_header, colWidths=[400, 140])
    t_lab.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0.5),
        ('TOPPADDING', (0, 0), (-1, -1), 0.5),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(t_lab)

    bullets_lab = [
        "Built responsive React frontend components, improving UI performance and cross-device compatibility.",
        "Integrated Express backend controllers with MongoDB databases to implement secure web endpoints.",
        "Authored comprehensive API documentation and local setup guides to accelerate developer onboarding.",
        "Participated in peer code reviews, troubleshooting codebase bugs and improving code readability."
    ]
    for b in bullets_lab:
        story.append(Paragraph(f"&bull; {b}", bullet_style))
    story.append(Spacer(1, 1))

    # DecodeLabs
    decodelabs_header = [
        [
            Paragraph("<b>DecodeLabs</b> – Web Development Intern", body_style),
            Paragraph("Summer 2025", ParagraphStyle(
                'RightText', parent=body_style, alignment=TA_RIGHT))
        ]
    ]
    t_dec = Table(decodelabs_header, colWidths=[400, 140])
    t_dec.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0.5),
        ('TOPPADDING', (0, 0), (-1, -1), 0.5),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(t_dec)

    bullets_dec = [
        "Created responsive, mobile-friendly layouts and modular site sections using React and Tailwind CSS.",
        "Connected React client components to Node.js REST API endpoints to deliver dynamic data updates.",
        "Utilized Git and GitHub for version control, branching, and pull requests in a team environment.",
        "Conducted local component and route testing to ensure strict alignment with visual design mockups."
    ]
    for b in bullets_dec:
        story.append(Paragraph(f"&bull; {b}", bullet_style))
    story.append(Spacer(1, 1))

    # --- 5. FEATURED PROJECTS ---
    add_section_header("Featured Projects")

    # Portfolio OS
    story.append(Paragraph(
        "<b>Rishabh Portfolio OS</b> (Next.js, React, TypeScript, "
        "Tailwind CSS, Git) | <i><a href=\"https://github.com/"
        "rishabhjain071130-glitch/rishabh-jain\"><font color=\"#1B365D\">"
        "<u>github.com/rishabhjain071130-glitch/rishabh-jain</u>"
        "</font></a></i>",
        body_style
    ))
    story.append(Paragraph(
        "&bull; Designed terminal-themed developer portfolio featuring "
        "CLI console simulation and rate-limited forms.", bullet_style
    ))

    # EduPilot AI
    story.append(Paragraph(
        "<b>EduPilot AI</b> (Python, Streamlit, Gemini API, HTML, CSS, Git) | "
        "<i><a href=\"https://github.com/rishabhjain071130-glitch/"
        "EduPilot-AI\"><font color=\"#1B365D\"><u>"
        "github.com/rishabhjain071130-glitch/EduPilot-AI</u>"
        "</font></a></i>",
        body_style
    ))
    story.append(Paragraph(
        "&bull; Built AI career mentor mapping skill gaps and generating "
        "custom roadmaps using Gemini API prompts.", bullet_style
    ))

    # Digital Loan System
    story.append(Paragraph(
        "<b>Digital Loan System</b> (Node.js, Express.js, MongoDB, "
        "Mongoose, HTML, CSS) | <i><a href=\"https://github.com/"
        "rishabhjain071130-glitch/digitalloansystem\"><font color=\"#1B365D\">"
        "<u>github.com/rishabhjain071130-glitch/digitalloansystem</u>"
        "</font></a></i>",
        body_style
    ))
    story.append(Paragraph(
        "&bull; Developed Express/MongoDB loan manager with session control, "
        "ledger logs, and automated notifications.", bullet_style
    ))
    story.append(Spacer(1, 1))

    # --- 6. CERTIFICATIONS ---
    add_section_header("Certifications")
    certs = [
        "<b>Oracle Database SQL Certified Associate</b> – SQL queries, "
        "database indexing, and relational models (Dec 2025)",
        "<b>Microsoft Agentic AI Certification</b> – Development and deployment "
        "of autonomous LLM agent systems (June 2026)",
        "<b>Microsoft Web Technology Specialization</b> – Standard-compliant "
        "frontend and backend implementations (June 2026)",
        "<b>Palo Alto Networks Network Security Gateway</b> – Network defense "
        "and firewall rule configurations (June 2026)",
        "<b>IBM Machine Learning Roadmap Certification</b> – Core ML models, "
        "pipelines, and data architectures (Jan 2026)",
        "<b>Google Generative AI Course Completion</b> – Prompt engineering, "
        "LLM tuning, and AI safety practices (May 2026)",
        "<b>Internshala Cyber Security with AI Training</b> – Threat detection "
        "and AI-assisted security response (Jan 2026)"
    ]
    for c in certs:
        story.append(Paragraph(f"&bull; {c}", bullet_style))
    story.append(Spacer(1, 1))

    # --- 7. ACHIEVEMENTS ---
    add_section_header("Achievements")
    achievements = [
        "Earned 10+ professional technical credentials from Microsoft, Oracle, Palo Alto Networks, and IBM.",
        "Deployed 3 production-grade open-source applications addressing real-world utility and security.",
        "Maintained a strong academic record and verified internship credentials in GLA University cohort."
    ]
    for a in achievements:
        story.append(Paragraph(f"&bull; {a}", bullet_style))

    doc.build(story)


if __name__ == '__main__':
    main()
